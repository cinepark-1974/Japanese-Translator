"""
BLUE JEANS PICTURES — Japanese-Translator v1.0
한국어 시나리오 → 일본어 번역 (5-Stage Market Adaptation Pipeline)
Powered by Anthropic Claude API

Pipeline:
  Stage 1: Raw Translation (Sonnet) — 충실한 직역 + 캐릭터/통화/문화 매핑
  Stage 2: Format Conversion (Rule-based) — 한국 포맷 → 일본 〇柱 포맷
  Stage 3: Voice Rewrite (Opus) — 번역체 제거, 일본 시나리오 문체
  Stage 4: Dialogue Polish (Opus) — 경어 설계, 대사 현지화
  Stage 5: QA Check (Sonnet) — 포맷/경어/문화코드/스토리 검증 리포트
"""

import streamlit as st
import anthropic
import re
import io
import csv
import json
import time

from prompt import (
    STYLE_PRESETS,
    KEIGO_TONE_TAGS,
    MODEL_POLICY,
    STAGE_2_FORMAT_RULES,
    CURRENCY_RULE,
    CULTURAL_CODE_MAP,
    build_stage1_prompt,
    build_stage3_prompt,
    build_stage4_prompt,
    build_stage5_prompt,
)

# ─────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="Japanese-Translator v1.0 | BLUE JEANS PICTURES",
    page_icon="🎌",
    layout="wide",
)

# ─────────────────────────────────────────────
# CUSTOM CSS
# ─────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700;900&display=swap');

.stApp { background-color: #F7F7F5 !important; }

.main-header {
    text-align: center;
    padding: 2.5rem 0 1rem 0;
}
.main-header .brand-name {
    font-size: 0.85rem; color: #191970;
    letter-spacing: 0.35em; font-weight: 600; margin-bottom: 0.3rem;
}
.main-header h1 {
    font-family: 'Playfair Display', serif;
    font-size: 2.8rem; font-weight: 900; color: #191970;
    margin: 0.2rem 0 0; letter-spacing: 0.02em;
    display: inline-block; border-bottom: 4px solid #f5c842;
    padding-bottom: 0.3rem;
}
.main-header .tagline {
    font-size: 0.78rem; color: #999;
    letter-spacing: 0.3em; margin-top: 0.7rem; font-weight: 400;
}
.main-header .version-badge {
    display: inline-block; background: #191970; color: #f5c842;
    font-size: 0.7rem; padding: 0.15rem 0.6rem; border-radius: 10px;
    margin-top: 0.5rem; letter-spacing: 0.1em; font-weight: 600;
}

.section-header {
    background: #f5c842; color: #191970;
    padding: 0.5rem 1.1rem; border-radius: 6px;
    font-weight: 700; font-size: 0.95rem;
    margin: 1.5rem 0 0.8rem 0; letter-spacing: 0.03em;
}

.stage-badge {
    display: inline-block; padding: 0.3rem 0.8rem;
    border-radius: 15px; font-size: 0.8rem;
    font-weight: 700; margin: 0.3rem 0.2rem; letter-spacing: 0.02em;
}
.stage-active { background: #191970; color: #f5c842; }
.stage-done { background: #2ecc71; color: #fff; }
.stage-pending { background: #ddd; color: #999; }

.result-box {
    background: #fff; border: 1px solid #ddd; border-radius: 8px;
    padding: 1.2rem; font-family: 'Yu Gothic', 'Hiragino Sans', sans-serif;
    font-size: 0.88rem; line-height: 1.7;
    white-space: pre-wrap; max-height: 600px; overflow-y: auto;
}

.qa-box {
    background: #FFFEF5; border: 2px solid #f5c842; border-radius: 8px;
    padding: 1.2rem; font-family: 'Courier New', monospace;
    font-size: 0.85rem; line-height: 1.6;
    white-space: pre-wrap; max-height: 500px; overflow-y: auto;
}

.page-chip {
    display: inline-block; background: #191970; color: #f5c842;
    padding: 0.2rem 0.7rem; border-radius: 12px;
    font-size: 0.8rem; font-weight: 600; margin-bottom: 0.5rem;
}

.char-table {
    width: 100%; border-collapse: collapse; margin: 0.5rem 0; font-size: 0.85rem;
}
.char-table th {
    background: #191970; color: #f5c842;
    padding: 0.45rem 0.8rem; text-align: left; font-weight: 600;
}
.char-table td {
    padding: 0.4rem 0.8rem; border-bottom: 1px solid #e8e8e0;
}
.char-table tr:nth-child(even) td { background: #EEEEF6; }

.pipeline-info {
    background: #EEEEF6; border-left: 4px solid #191970;
    padding: 0.8rem 1rem; border-radius: 0 6px 6px 0;
    font-size: 0.85rem; margin: 0.5rem 0; color: #333;
}

.progress-text {
    text-align: center; color: #666; font-size: 0.85rem; padding: 0.5rem 0;
}

.footer {
    text-align: center; color: #bbb; font-size: 0.72rem;
    margin-top: 2.5rem; padding: 1rem 0;
    border-top: 1px solid #ddd; letter-spacing: 0.08em;
}

.stTextInput > div > div > input,
.stTextArea > div > div > textarea,
.stSelectbox > div > div { background-color: #fff !important; border-color: #ddd !important; }
div[data-testid="stFileUploader"] { background-color: #fff !important; border-radius: 8px; }
.stExpander { background-color: #fff !important; border-color: #ddd !important; border-radius: 8px !important; }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────
MAX_CHARS_PER_PAGE = 8000
VERSION = "1.0"


# ─────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────

def parse_character_map(uploaded_file) -> tuple:
    """Parse character name mapping from CSV or TXT file.
    Supports optional keigo tag and note columns:
    한국이름,일본이름,경어태그,호칭메모
    """
    name = uploaded_file.name.lower()
    content = uploaded_file.read().decode("utf-8", errors="replace")
    uploaded_file.seek(0)

    char_map = {}
    char_tones = {}
    skip_headers = {"한국이름", "한국명", "korean", "name", "이름"}

    if name.endswith(".csv"):
        reader = csv.reader(io.StringIO(content))
        for row in reader:
            if len(row) >= 2:
                ko = row[0].strip()
                jp = row[1].strip()
                if ko and jp and ko.lower() not in skip_headers:
                    char_map[ko] = jp
                    if len(row) >= 3:
                        tone = row[2].strip().lower()
                        if tone in KEIGO_TONE_TAGS:
                            char_tones[jp] = tone
    else:
        for line in content.strip().split("\n"):
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            for sep in ["→", "->", "=>", "=", ",", ":", "\t"]:
                if sep in line:
                    parts = line.split(sep)
                    ko = parts[0].strip()
                    jp = parts[1].strip() if len(parts) > 1 else ""
                    if ko and jp and ko.lower() not in skip_headers:
                        char_map[ko] = jp
                        if len(parts) >= 3:
                            tone = parts[2].strip().lower()
                            if tone in KEIGO_TONE_TAGS:
                                char_tones[jp] = tone
                    break

    return char_map, char_tones


def read_uploaded_file(uploaded_file) -> str:
    """Read text from uploaded .txt, .pdf, or .docx file."""
    name = uploaded_file.name.lower()

    if name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="replace")

    elif name.endswith(".pdf"):
        try:
            import pymupdf
            pdf_bytes = uploaded_file.read()
            doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
            return "\n".join(text_parts)
        except ImportError:
            st.error("PDF 처리를 위해 pymupdf가 필요합니다.")
            return ""

    elif name.endswith(".docx"):
        try:
            from docx import Document
            docx_bytes = io.BytesIO(uploaded_file.read())
            doc = Document(docx_bytes)
            return "\n".join([p.text for p in doc.paragraphs])
        except ImportError:
            st.error("DOCX 처리를 위해 python-docx가 필요합니다.")
            return ""

    else:
        st.error("지원하지 않는 파일 형식입니다. (.txt / .pdf / .docx)")
        return ""


def split_into_pages(text: str, max_chars: int = MAX_CHARS_PER_PAGE) -> list:
    """Split text into pages, trying to break at scene boundaries."""
    if len(text) <= max_chars:
        return [text]

    pages = []
    remaining = text

    while remaining:
        if len(remaining) <= max_chars:
            pages.append(remaining)
            break

        chunk = remaining[:max_chars]
        # 한국어 씬 패턴 + 일본어 〇 패턴
        scene_pattern = r'\n\s*(?:S\s*#?\s*\d+|씬\s*#?\s*\d+|SCENE\s*\d+|#\s*\d+[\.\)]|INT\.|EXT\.|〇)'
        matches = list(re.finditer(scene_pattern, chunk))

        if matches:
            cut_point = matches[-1].start()
            if cut_point > max_chars * 0.3:
                pages.append(remaining[:cut_point].rstrip())
                remaining = remaining[cut_point:].lstrip("\n")
                continue

        last_break = chunk.rfind("\n\n")
        if last_break > max_chars * 0.3:
            pages.append(remaining[:last_break].rstrip())
            remaining = remaining[last_break:].lstrip("\n")
        else:
            pages.append(chunk)
            remaining = remaining[max_chars:]

    return pages


# ─────────────────────────────────────────────
# STAGE 2: FORMAT CONVERSION (Rule-based)
# ─────────────────────────────────────────────

def apply_format_conversion(text: str) -> str:
    """Apply rule-based format conversion: Korean screenplay → Japanese 横書き format."""
    result = text

    # ── Scene heading: number prefix → 〇 ──
    # Patterns: "1. INT. ..." / "1. EXT. ..." / "S#1. INT. ..." / "씬1. ..."
    result = re.sub(
        r'^(?:\d+\.?\s*|S\s*#?\s*\d+\.?\s*|씬\s*\d+\.?\s*)'
        r'(?:INT\.\s*/??\s*EXT\.\s*|EXT\.\s*/??\s*INT\.\s*|INT\.\s*|EXT\.\s*)?',
        '〇',
        result,
        flags=re.MULTILINE | re.IGNORECASE
    )

    # ── Scene heading: dash time → parenthetical time ──
    # "〇場所 — 時間" → "〇場所（時間）"
    result = re.sub(
        r'^(〇.+?)\s*[—–\-]\s*(.+)$',
        lambda m: f"{m.group(1).rstrip()}（{m.group(2).strip()}）",
        result,
        flags=re.MULTILINE
    )

    # ── Korean time words → Japanese ──
    time_map = {
        "새벽": "夜明け", "아침": "朝", "오전": "午前", "낮": "昼",
        "오후": "午後", "저녁": "夕方", "밤": "夜", "심야": "深夜",
        "늦은 오후": "夕方", "늦은 밤": "深夜",
    }
    for ko_time, jp_time in time_map.items():
        result = result.replace(f"（{ko_time}）", f"（{jp_time}）")

    # ── Korean direction markers → Japanese ──
    marker_map = {
        r'\(N\)': '（ナレーション）',
        r'\(나레이션\)': '（ナレーション）',
        r'\(V\.O\.?\)': '（声のみ）',
        r'\(O\.S\.?\)': '（OFF）',
        r'\(소리\)': '（OFF）',
        r'\(독백\)': '（ナレーション）',
        r'\(전화\)': '（電話）',
        r'\(계속\)': '（続き）',
        r'\(회상\)': '（回想）',
        r'\(몽타주\)': '（モンタージュ）',
        r'\(타이틀\)': 'タイトル。',
        r'\(자막\)': '字幕。',
    }
    for pattern, replacement in marker_map.items():
        result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)

    # ── Transitions ──
    result = re.sub(r'CUT\s+TO\.?', 'カットバック。', result, flags=re.IGNORECASE)
    result = re.sub(r'INSERT\s*>', 'インサート＞', result, flags=re.IGNORECASE)
    result = re.sub(r'FADE\s+IN\s*:', 'フェイドイン。', result, flags=re.IGNORECASE)
    result = re.sub(r'FADE\s+OUT\.?', 'フェイドアウト。', result, flags=re.IGNORECASE)
    result = re.sub(r'SMASH\s+CUT\s*:', 'スマッシュカット。', result, flags=re.IGNORECASE)

    # ── V.O. in character cues ──
    result = re.sub(r'\(V\.O\.?\)', '（声のみ）', result)
    result = re.sub(r'\(O\.S\.?\)', '（OFF）', result)

    return result


# ─────────────────────────────────────────────
# API CALL FUNCTION
# ─────────────────────────────────────────────

def call_api(client, text: str, system_prompt: str, model_id: str,
             max_tokens: int = 8000, page_info: str = "") -> str:
    """Call Claude API with streaming to prevent timeout."""
    full_system = system_prompt
    if page_info:
        full_system += f"\n\n[Internal context: {page_info}. Do NOT include this in output.]"

    collected = []
    with client.messages.stream(
        model=model_id,
        max_tokens=max_tokens,
        system=full_system,
        messages=[{"role": "user", "content": text}]
    ) as stream:
        for text_chunk in stream.text_stream:
            collected.append(text_chunk)

    return "".join(collected)


def run_stage_on_pages(client, pages: list, system_prompt: str,
                       model_id: str, stage_name: str,
                       progress_bar, status_area) -> list:
    """Run an API-based stage on multiple pages with progress tracking."""
    results = []
    total = len(pages)

    for idx, page in enumerate(pages):
        page_num = idx + 1
        status_area.markdown(
            f'<div class="progress-text">🔄 {stage_name} — 페이지 {page_num}/{total} 처리 중... (모델: {model_id})</div>',
            unsafe_allow_html=True
        )

        try:
            result = call_api(
                client, page, system_prompt, model_id,
                page_info=f"Page {page_num} of {total}. Maintain consistency."
            )
            results.append(result)
        except anthropic.APIError as e:
            error_msg = f"❌ API 오류 ({stage_name}, 페이지 {page_num}): {e}"
            st.error(error_msg)
            st.session_state["last_error"] = error_msg
            return None
        except Exception as e:
            error_msg = f"❌ 오류 ({stage_name}, 페이지 {page_num}): {type(e).__name__}: {e}"
            st.error(error_msg)
            st.session_state["last_error"] = error_msg
            return None

        progress_bar.progress(page_num / total)

    return results


# ─────────────────────────────────────────────
# DOCX GENERATION (横書き A4)
# ─────────────────────────────────────────────

def generate_docx(text: str) -> bytes:
    """Generate formatted Japanese screenplay DOCX (横書き A4)."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Page Setup: A4 横書き ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ── Styles ──
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Yu Gothic'
    style_normal.font.size = Pt(10.5)
    style_normal.paragraph_format.space_after = Pt(0)
    style_normal.paragraph_format.space_before = Pt(0)
    style_normal.paragraph_format.line_spacing = 1.5

    # Scene heading (柱)
    style_scene = doc.styles.add_style('Hashira', 1)
    style_scene.font.name = 'Yu Gothic'
    style_scene.font.size = Pt(10.5)
    style_scene.font.bold = True
    style_scene.paragraph_format.space_before = Pt(24)
    style_scene.paragraph_format.space_after = Pt(6)
    style_scene.paragraph_format.line_spacing = 1.5

    # Action (ト書き) — 3字下げ
    style_action = doc.styles.add_style('Togaki', 1)
    style_action.font.name = 'Yu Gothic'
    style_action.font.size = Pt(10.5)
    style_action.paragraph_format.left_indent = Cm(1.0)  # ~3字下げ
    style_action.paragraph_format.space_before = Pt(3)
    style_action.paragraph_format.space_after = Pt(0)
    style_action.paragraph_format.line_spacing = 1.5

    # Character name (人物名)
    style_char = doc.styles.add_style('Jinmei', 1)
    style_char.font.name = 'Yu Gothic'
    style_char.font.size = Pt(10.5)
    style_char.font.bold = True
    style_char.paragraph_format.space_before = Pt(6)
    style_char.paragraph_format.space_after = Pt(0)
    style_char.paragraph_format.line_spacing = 1.5

    # Dialogue (セリフ)
    style_dialog = doc.styles.add_style('Serifu', 1)
    style_dialog.font.name = 'Yu Gothic'
    style_dialog.font.size = Pt(10.5)
    style_dialog.paragraph_format.space_before = Pt(0)
    style_dialog.paragraph_format.space_after = Pt(0)
    style_dialog.paragraph_format.line_spacing = 1.5

    # Transition
    style_trans = doc.styles.add_style('Tenkan', 1)
    style_trans.font.name = 'Yu Gothic'
    style_trans.font.size = Pt(10.5)
    style_trans.paragraph_format.space_before = Pt(6)
    style_trans.paragraph_format.space_after = Pt(6)
    style_trans.paragraph_format.line_spacing = 1.5

    # ── Parse and format ──
    HASHIRA_RE = re.compile(r'^〇')
    TRANSITION_RE = re.compile(
        r'^(カットバック|インサート|フェイドイン|フェイドアウト|スマッシュカット|タイトル|字幕|モンタージュ)',
    )
    # セリフ: name「...」 or name（声のみ）「...」
    SERIFU_RE = re.compile(r'^(.+?)(?:（[^）]*）)?\s*「')

    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # 柱
        if HASHIRA_RE.match(stripped):
            doc.add_paragraph(stripped, style='Hashira')
            i += 1
            continue

        # 転換
        if TRANSITION_RE.match(stripped):
            doc.add_paragraph(stripped, style='Tenkan')
            i += 1
            continue

        # セリフ (「」を含む行)
        if '「' in stripped:
            # 人物名 + セリフ が同一行
            m = SERIFU_RE.match(stripped)
            if m:
                doc.add_paragraph(stripped, style='Serifu')
                i += 1
                continue

        # ト書き (default)
        doc.add_paragraph(stripped, style='Togaki')
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ═══════════════════════════════════════════════════
# UI
# ═══════════════════════════════════════════════════

# ── Header ──
st.markdown(f"""
<div class="main-header">
    <div class="brand-name">B L U E &nbsp; J E A N S &nbsp; P I C T U R E S</div>
    <h1>JAPANESE-TRANSLATOR</h1>
    <div class="tagline">Y O U N G &nbsp; · &nbsp; V I N T A G E &nbsp; · &nbsp; F R E E &nbsp; · &nbsp; I N N O V A T I V E</div>
    <div class="version-badge">v{VERSION} — 5-Stage Market Adaptation Pipeline</div>
</div>
""", unsafe_allow_html=True)

st.markdown("---")

# ── API Key ──
api_key = st.secrets.get("ANTHROPIC_API_KEY", "")
if api_key:
    st.success("🔑 API Key 연결됨 (Secrets)")
else:
    api_key = st.text_input(
        "🔑 Anthropic API Key",
        type="password",
        help="Claude API 키를 입력하세요. (sk-ant-...)"
    )

# ═══════════════════════════════════════════════════
# SETTINGS PANEL
# ═══════════════════════════════════════════════════

st.markdown('<div class="section-header">⚙️ SETTINGS — 번역 설정</div>', unsafe_allow_html=True)

# ── Style ──
st.markdown("**장르 스타일**")
style_choice = st.selectbox(
    "장르/스타일 프리셋:",
    list(STYLE_PRESETS.keys()),
    index=0,
    label_visibility="collapsed",
)
selected_style = STYLE_PRESETS[style_choice]
st.caption(f"📌 {selected_style['desc']}")

# ── Custom Instructions ──
custom_instructions = st.text_area(
    "✏️ 추가 번역 지시사항 (선택)",
    height=80,
    placeholder="예: 특정 용어는 이렇게 번역해줘 / 경어 레벨을 이렇게 조정해줘...",
)

# ── Pipeline Info ──
st.markdown(
    '<div class="pipeline-info"><strong>5-Stage Market Adaptation Pipeline:</strong><br>'
    'Stage 1: Raw Translation → Sonnet (직역 + 캐릭터/통화/문화 매핑)<br>'
    'Stage 2: Format Conversion → 규칙 기반 (무료, 〇柱 포맷 변환)<br>'
    'Stage 3: Voice Rewrite → Opus (번역체 제거, 일본 시나리오 문체)<br>'
    'Stage 4: Dialogue Polish → Opus (경어 설계, 대사 현지화)<br>'
    'Stage 5: QA Check → Sonnet (포맷/경어/문화코드/스토리 검증)<br>'
    '<br>💡 각 단계별로 독립 실행 · 결과 저장 · 이어서 진행 가능</div>',
    unsafe_allow_html=True
)


# ═══════════════════════════════════════════════════
# CHARACTER MAP
# ═══════════════════════════════════════════════════

st.markdown('<div class="section-header">👤 CHARACTER MAP — 인물표</div>', unsafe_allow_html=True)

st.info("💡 CSV/TXT 파일을 업로드하세요. 3번째 열에 경어 태그(keigo/teinei/tameguchi/kenson/ibar) 추가 가능!")

char_map_file = st.file_uploader(
    "인물표 파일 업로드",
    type=["csv", "txt"],
    help="CSV: 한국이름,일본이름,경어태그 | TXT: 한국이름 → 일본이름 → 경어태그",
    key="char_map_upload"
)

char_map = {}
char_tones = {}

if char_map_file:
    char_map, char_tones = parse_character_map(char_map_file)
    if char_map:
        st.success(f"✅ {len(char_map)}명 매핑 로드 · {len(char_tones)}명 경어 태그 설정됨")
        rows = []
        for ko, jp in char_map.items():
            tone = char_tones.get(jp, "—")
            tone_label = KEIGO_TONE_TAGS[tone]["label"] if tone in KEIGO_TONE_TAGS else "—"
            rows.append(f"<tr><td>{ko}</td><td>→</td><td><strong>{jp}</strong></td><td>{tone_label}</td></tr>")
        st.markdown(f"""
        <table class="char-table">
            <tr><th>한국이름</th><th></th><th>日本語名</th><th>敬語レベル</th></tr>
            {"".join(rows)}
        </table>
        """, unsafe_allow_html=True)
    else:
        st.warning("⚠️ 인물 매핑을 읽을 수 없습니다.")

# Manual keigo assignment
if char_map and not char_tones:
    with st.expander("🎭 캐릭터별 경어 태그 수동 설정 (선택)"):
        st.caption("CSV 3번째 열 없이 여기서 직접 설정할 수 있어요.")
        for ko, jp in char_map.items():
            tone = st.selectbox(
                f"{jp}",
                ["—", "keigo", "teinei", "tameguchi", "kenson", "ibar"],
                index=0,
                key=f"tone_{jp}",
            )
            if tone != "—":
                char_tones[jp] = tone

with st.expander("📋 인물표 파일 예시 보기"):
    st.code("""# CSV 형식 — 4열: 한국이름, 일본이름, 경어태그, 호칭메모(선택)
한국이름,일본이름,경어태그,호칭메모
김지훈,中村智弘,teinei,松田には敬語
이수현,林美咲,teinei,智弘とは距離がある
박성태,杉山正道,ibar,タメ口+命令形
박 경위,松田警部補,tameguchi,職務的タメ口

# TXT 형식
김지훈 → 中村智弘 → teinei
박성태 → 杉山正道 → ibar""", language="text")


# ═══════════════════════════════════════════════════
# INPUT
# ═══════════════════════════════════════════════════

st.markdown('<div class="section-header">📥 INPUT — 시나리오 입력 (한국어)</div>', unsafe_allow_html=True)

project_title = st.text_input(
    "🎬 프로젝트 제목 (파일명에 사용됩니다)",
    value=st.session_state.get("project_title", ""),
    placeholder="예: 水姫, 물귀신, TAKEOFF...",
    help="다운로드 파일명이 Screenplay_제목_JP 형태로 생성됩니다.",
)
if project_title:
    st.session_state["project_title"] = project_title

input_method = st.radio(
    "입력 방식:",
    ["📎 파일 업로드", "📝 텍스트 붙여넣기"],
    horizontal=True,
)

source_text = ""

if input_method == "📎 파일 업로드":
    uploaded = st.file_uploader(
        "시나리오 파일을 업로드하세요",
        type=["txt", "pdf", "docx"],
        help=".txt / .pdf / .docx 파일 지원",
        key="screenplay_upload"
    )
    if uploaded:
        with st.spinner("파일 읽는 중..."):
            source_text = read_uploaded_file(uploaded)
        if source_text:
            import os
            raw_name = os.path.splitext(uploaded.name)[0]
            clean_title = re.sub(r'[\s_-]+', '_', raw_name).strip('_')
            st.session_state["project_title"] = clean_title
            st.success(f"✅ 파일 로드 완료 — {len(source_text):,}자")
            with st.expander("📄 원문 미리보기", expanded=False):
                st.text(source_text[:3000] + ("..." if len(source_text) > 3000 else ""))
else:
    if "paste_pages" not in st.session_state:
        st.session_state.paste_pages = 1

    col_add, col_remove = st.columns([1, 1])
    with col_add:
        if st.button("➕ 페이지 추가", use_container_width=True):
            st.session_state.paste_pages += 1
            st.rerun()
    with col_remove:
        if st.session_state.paste_pages > 1:
            if st.button("➖ 페이지 제거", use_container_width=True):
                st.session_state.paste_pages -= 1
                st.rerun()

    page_texts = []
    for i in range(st.session_state.paste_pages):
        st.markdown(f'<span class="page-chip">Page {i+1}</span>', unsafe_allow_html=True)
        txt = st.text_area(
            f"시나리오 텍스트 (페이지 {i+1})",
            height=250,
            key=f"paste_page_{i}",
            placeholder=f"페이지 {i+1}의 시나리오 텍스트를 붙여넣으세요...",
            label_visibility="collapsed"
        )
        page_texts.append(txt)

    source_text = "\n\n".join([t for t in page_texts if t.strip()])
    if source_text:
        st.caption(f"총 {len(source_text):,}자 입력됨")


# ═══════════════════════════════════════════════════
# STEP-BY-STEP PIPELINE
# ═══════════════════════════════════════════════════

st.markdown('<div class="section-header">🔄 PIPELINE — 단계별 실행</div>', unsafe_allow_html=True)

can_run = bool(api_key and source_text.strip())

if not api_key:
    st.warning("⬆️ API Key를 먼저 입력하세요.")
elif not source_text.strip():
    st.warning("⬆️ 시나리오 텍스트를 입력하세요.")

for key in ["stage_1_result", "stage_2_result", "stage_3_result", "stage_4_result", "stage_5_result"]:
    if key not in st.session_state:
        st.session_state[key] = None


def show_stage_result(stage_num: int, stage_name: str, result_key: str):
    result = st.session_state.get(result_key)
    if result:
        with st.expander(f"📄 Stage {stage_num} 결과 — {stage_name}", expanded=False):
            st.text(result[:5000] + ("..." if len(result) > 5000 else ""))
        st.download_button(
            f"💾 Stage {stage_num} 결과 저장 (TXT)",
            data=result.encode("utf-8"),
            file_name=f"stage_{stage_num}_{stage_name.lower().replace(' ', '_')}.txt",
            mime="text/plain",
            use_container_width=True,
            key=f"dl_stage_{stage_num}",
        )


def upload_previous_result(stage_num: int, prev_stage_name: str, result_key: str):
    prev_result = st.session_state.get(result_key)
    if prev_result:
        st.success(f"✅ 이전 단계 결과 있음 ({len(prev_result):,}자) — 자동 연결됩니다.")
        return prev_result
    else:
        st.info(f"💡 이전 단계({prev_stage_name}) 결과가 없으면 파일을 업로드하세요.")
        prev_file = st.file_uploader(
            f"Stage {stage_num - 1} 결과 파일 업로드",
            type=["txt"],
            key=f"prev_upload_{stage_num}",
        )
        if prev_file:
            text = prev_file.read().decode("utf-8", errors="replace")
            st.success(f"✅ 파일 로드 완료 — {len(text):,}자")
            return text
    return None


# ═══════════════════════════════════════════════════
# STAGE 1: Raw Translation
# ═══════════════════════════════════════════════════
st.markdown("---")
st.markdown("### ① Raw Translation (Sonnet)")
st.caption("한국어 → 일본어 직역. 캐릭터명/통화/문화코드 매핑 적용.")

if can_run:
    if st.button("▶️ Stage 1 실행", key="btn_stage1", use_container_width=True):
        client = anthropic.Anthropic(api_key=api_key)

        system_prompt = build_stage1_prompt(
            char_map=char_map,
            style_prompt=selected_style["prompt"],
            custom_instructions=custom_instructions,
        )
        model_id = MODEL_POLICY["stage_1"]["model"]
        pages = split_into_pages(source_text)

        progress_bar = st.progress(0)
        status_area = st.empty()

        results = run_stage_on_pages(
            client, pages, system_prompt, model_id,
            "Stage 1: Raw Translation", progress_bar, status_area
        )

        if results is not None:
            st.session_state["stage_1_result"] = "\n\n".join(results)
            status_area.markdown('<div class="progress-text">✅ Stage 1 완료!</div>', unsafe_allow_html=True)
            st.rerun()

show_stage_result(1, "Raw Translation", "stage_1_result")


# ═══════════════════════════════════════════════════
# STAGE 2: Format Conversion
# ═══════════════════════════════════════════════════
st.markdown("---")
st.markdown("### ② Format Conversion (규칙 기반)")
st.caption("한국 씬 헤더 → 일본 〇柱 포맷. API 호출 없음 (무료).")

stage_2_input = st.session_state.get("stage_1_result")
if stage_2_input:
    if st.button("▶️ Stage 2 실행", key="btn_stage2", use_container_width=True):
        st.session_state["stage_2_result"] = apply_format_conversion(stage_2_input)
        st.rerun()
elif st.session_state.get("stage_2_result") is None:
    st.caption("⏳ Stage 1을 먼저 완료하세요.")

show_stage_result(2, "Format", "stage_2_result")


# ═══════════════════════════════════════════════════
# STAGE 3: Voice Rewrite
# ═══════════════════════════════════════════════════
st.markdown("---")
st.markdown("### ③ Voice Rewrite (Opus)")
st.caption("번역체 제거, 일본 시나리오 문체로 리라이트. JP-1~JP-9 안티패턴 교정.")

stage_3_input = upload_previous_result(3, "Stage 2 Format", "stage_2_result")
if stage_3_input and api_key:
    if st.button("▶️ Stage 3 실행", key="btn_stage3", use_container_width=True):
        client = anthropic.Anthropic(api_key=api_key)

        system_prompt = build_stage3_prompt(
            char_map=char_map,
            char_tones=char_tones,
            style_prompt=selected_style["prompt"],
            custom_instructions=custom_instructions,
        )
        model_id = MODEL_POLICY["stage_3"]["model"]
        pages = split_into_pages(stage_3_input)

        progress_bar = st.progress(0)
        status_area = st.empty()

        results = run_stage_on_pages(
            client, pages, system_prompt, model_id,
            "Stage 3: Voice Rewrite", progress_bar, status_area
        )

        if results is not None:
            st.session_state["stage_3_result"] = "\n\n".join(results)
            status_area.markdown('<div class="progress-text">✅ Stage 3 완료!</div>', unsafe_allow_html=True)
            st.rerun()
elif st.session_state.get("stage_3_result") is None:
    st.caption("⏳ Stage 2를 먼저 완료하세요.")

show_stage_result(3, "Voice Rewrite", "stage_3_result")


# ═══════════════════════════════════════════════════
# STAGE 4: Dialogue Polish
# ═══════════════════════════════════════════════════
st.markdown("---")
st.markdown("### ④ Dialogue Polish (Opus)")
st.caption("경어 설계 적용, 대사 전문 폴리시. DP-1~DP-6 원칙.")

stage_4_input = upload_previous_result(4, "Stage 3 Voice Rewrite", "stage_3_result")
if stage_4_input and api_key:
    if st.button("▶️ Stage 4 실행", key="btn_stage4", use_container_width=True):
        client = anthropic.Anthropic(api_key=api_key)

        system_prompt = build_stage4_prompt(
            char_map=char_map,
            char_tones=char_tones,
            style_prompt=selected_style["prompt"],
            custom_instructions=custom_instructions,
        )
        model_id = MODEL_POLICY["stage_4"]["model"]
        pages = split_into_pages(stage_4_input)

        progress_bar = st.progress(0)
        status_area = st.empty()

        results = run_stage_on_pages(
            client, pages, system_prompt, model_id,
            "Stage 4: Dialogue Polish", progress_bar, status_area
        )

        if results is not None:
            st.session_state["stage_4_result"] = "\n\n".join(results)
            status_area.markdown('<div class="progress-text">✅ Stage 4 완료!</div>', unsafe_allow_html=True)
            st.rerun()
elif st.session_state.get("stage_4_result") is None:
    st.caption("⏳ Stage 3를 먼저 완료하세요.")

show_stage_result(4, "Dialogue Polish", "stage_4_result")


# ═══════════════════════════════════════════════════
# STAGE 5: QA Check
# ═══════════════════════════════════════════════════
st.markdown("---")
st.markdown("### ⑤ QA Check (Sonnet)")
st.caption("최종 품질 검증. 포맷/경어/문화코드/스토리 체크리스트.")

stage_5_input = upload_previous_result(5, "Stage 4 Dialogue Polish", "stage_4_result")
if stage_5_input and api_key:
    if st.button("▶️ Stage 5 실행", key="btn_stage5", use_container_width=True):
        client = anthropic.Anthropic(api_key=api_key)

        system_prompt = build_stage5_prompt()
        model_id = MODEL_POLICY["stage_5"]["model"]

        status_area = st.empty()
        status_area.markdown(
            '<div class="progress-text">🔍 QA 검증 중...</div>',
            unsafe_allow_html=True
        )

        try:
            qa_input = stage_5_input
            if len(qa_input) > 30000:
                qa_input = stage_5_input[:15000] + "\n\n[...中略...]\n\n" + stage_5_input[-15000:]

            qa_report = call_api(
                client, qa_input, system_prompt,
                model_id, max_tokens=4000
            )
            st.session_state["stage_5_result"] = qa_report
            status_area.markdown('<div class="progress-text">✅ Stage 5 완료!</div>', unsafe_allow_html=True)
            st.rerun()
        except Exception as e:
            st.error(f"❌ QA 오류: {e}")
elif st.session_state.get("stage_5_result") is None:
    st.caption("⏳ Stage 4를 먼저 완료하세요.")

if st.session_state.get("stage_5_result"):
    st.markdown("**🔍 QA Report**")
    st.markdown(f'<div class="qa-box">{st.session_state["stage_5_result"]}</div>', unsafe_allow_html=True)
    st.download_button(
        "💾 QA Report 저장 (TXT)",
        data=st.session_state["stage_5_result"].encode("utf-8"),
        file_name="qa_report_jp.txt",
        mime="text/plain",
        use_container_width=True,
        key="dl_qa",
    )


# ═══════════════════════════════════════════════════
# FINAL OUTPUT — DOCX
# ═══════════════════════════════════════════════════

final_result = (
    st.session_state.get("stage_4_result")
    or st.session_state.get("stage_3_result")
    or st.session_state.get("stage_2_result")
    or st.session_state.get("stage_1_result")
)

if final_result:
    st.markdown("---")
    st.markdown('<div class="section-header">📤 FINAL OUTPUT — 최종 다운로드</div>', unsafe_allow_html=True)

    title_slug = st.session_state.get("project_title", "").strip()
    if title_slug:
        title_slug = re.sub(r'[^\w\s\-]', '', title_slug).strip()
        title_slug = re.sub(r'[\s]+', '_', title_slug)
        base_filename = f"Screenplay_{title_slug}_JP"
    else:
        base_filename = "Screenplay_JP"

    if st.session_state.get("stage_4_result"):
        st.caption("✅ Stage 4 (Dialogue Polish) 결과 기준")
    elif st.session_state.get("stage_3_result"):
        st.caption("⚠️ Stage 3 (Voice Rewrite) 결과 기준 — Stage 4 미완료")
    elif st.session_state.get("stage_2_result"):
        st.caption("⚠️ Stage 2 (Format) 결과 기준 — Stage 3~4 미완료")
    else:
        st.caption("⚠️ Stage 1 (Raw Translation) 결과 기준 — 추가 폴리시 권장")

    col1, col2 = st.columns(2)

    with col1:
        st.download_button(
            "📥 TXT 다운로드",
            data=final_result.encode("utf-8"),
            file_name=f"{base_filename}.txt",
            mime="text/plain",
            use_container_width=True,
            key="dl_final_txt",
        )

    with col2:
        try:
            docx_bytes = generate_docx(final_result)
            st.download_button(
                "📥 DOCX 다운로드 (横書き A4)",
                data=docx_bytes,
                file_name=f"{base_filename}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="dl_final_docx",
            )
        except Exception as e:
            st.error(f"DOCX 생성 오류: {e}")

    st.markdown("")
    if st.button("🗑️ 전체 초기화 (새 프로젝트)", use_container_width=True):
        for key in ["stage_1_result", "stage_2_result", "stage_3_result",
                     "stage_4_result", "stage_5_result", "last_error"]:
            if key in st.session_state:
                del st.session_state[key]
        st.rerun()

if "last_error" in st.session_state:
    st.error(st.session_state["last_error"])


# ── Footer ──
st.markdown(f"""
<div class="footer">
    BLUE JEANS PICTURES — Japanese-Translator v{VERSION}<br>
    5-Stage Market Adaptation Pipeline · Powered by Anthropic Claude API
</div>
""", unsafe_allow_html=True)
